import os
import fitz  # PyMuPDF
from paddleocr import PaddleOCR
from PIL import Image
from docx import Document
from docx.shared import Pt, Mm
from docx.oxml.ns import qn
from deep_translator import GoogleTranslator  # For translation

# Initialize PaddleOCR for Chinese text
ocr = PaddleOCR(use_angle_cls=True, lang='ch')

# Initialize the Google Translator for deep-translator
translator = GoogleTranslator(source='zh-CN', target='vi')

# Path to folder containing PDF files
pdf_folder = 'input'
output_folder = 'output'

# Ensure output folder exists
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Create separate folders for images and Word files inside the output folder
images_folder = os.path.join(output_folder, 'images')
word_files_folder = os.path.join(output_folder, 'word_files')

# Ensure these subfolders exist
if not os.path.exists(images_folder):
    os.makedirs(images_folder)

if not os.path.exists(word_files_folder):
    os.makedirs(word_files_folder)

# A4 dimensions in points
A4_WIDTH_POINTS = 595.28  # A4 width in points
A4_HEIGHT_POINTS = 842.36  # A4 height in points

# Function to convert PDF page to high-resolution image (600 DPI)
def pdf_page_to_high_res_image(pdf_document, page_number):
    page = pdf_document.load_page(page_number)  # Load the page
    zoom = 20  # Set zoom factor for higher resolution (4x gives 600 DPI)
    mat = fitz.Matrix(zoom, zoom)  # Create matrix for zoom
    pix = page.get_pixmap(matrix=mat)  # Get the pixmap with the desired zoom
    image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)  # Convert to PIL Image
    return image, pix.width, pix.height  # Return image and dimensions

# Function to add text to Word with positional approximation and translation
def add_text_with_position(doc, text, bbox, image_width, image_height):
    # bbox: [[x0, y0], [x1, y1], [x2, y2], [x3, y3]] - quadrilateral coordinates of the text box in pixels
    # Extract top-left point from the quadrilateral
    x0, y0 = bbox[0]  # Top-left corner

    # Calculate scaling factors
    x_scale = A4_WIDTH_POINTS / image_width
    y_scale = A4_HEIGHT_POINTS / image_height

    # Scale the coordinates according to the canvas size
    x0_scaled = x0 * x_scale
    y0_scaled = y0 * y_scale

    # Translate the text to Vietnamese using deep-translator
    try:
        translated_text = translator.translate(text)
    except Exception as e:
        print(f"Error in translation: {e}")
        translated_text = text  # Fall back to the original text if translation fails

    # Create a paragraph with the translated text
    paragraph = doc.add_paragraph()

    # Add a run to the paragraph
    run = paragraph.add_run(translated_text)

    # Approximate positioning using indentation
    paragraph.paragraph_format.left_indent = Pt(x0_scaled)  # Indentation based on scaled x0

    # Set font size and other styles (optional)
    run.font.size = Pt(10)  # Adjust the font size as needed
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')  # Set font for Chinese text

# Loop through each PDF file in the folder
for filename in os.listdir(pdf_folder):
    if filename.endswith('.pdf'):
        pdf_path = os.path.join(pdf_folder, filename)
        print(f"Processing {pdf_path}")

        # Open the PDF document
        pdf_document = fitz.open(pdf_path)

        # Create a Word document for this PDF
        doc = Document()
        doc.add_heading(f'OCR Results for {filename} (Translated to Vietnamese)', 0)

        # Loop through all the pages in the PDF
        for page_num in range(pdf_document.page_count):
            # Convert PDF page to high-resolution image (600 DPI)
            image, image_width, image_height = pdf_page_to_high_res_image(pdf_document, page_num)

            # Save the image (inside the 'images' folder)
            image_filename = f"{filename}_page_{page_num + 1}_600dpi.png"
            image_path = os.path.join(images_folder, image_filename)
            image.save(image_path, 'PNG')

            # Perform OCR on the image
            result = ocr.ocr(image_path, cls=True)

            # Add page heading to Word document
            doc.add_heading(f'Page {page_num + 1}', level=1)

            # Add OCR results with positional information and translation
            for line in result:
                for word_info in line:
                    text = word_info[1][0]  # The recognized text
                    bbox = word_info[0]  # The bounding box coordinates
                    add_text_with_position(doc, text, bbox, image_width, image_height)

        # Save the Word document (inside the 'word_files' folder)
        word_output_path = os.path.join(word_files_folder, f"{filename}_ocr_results_translated.docx")
        doc.save(word_output_path)
        print(f"Translated OCR results saved to {word_output_path}")

        # Close the PDF document
        pdf_document.close()
